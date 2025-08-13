import os
import asyncio
import logging
import time
from pathlib import Path
from typing import Optional

from aiogram import Bot, Dispatcher, F, Router
from aiogram.filters import CommandStart
from aiogram.types import (
    Message, CallbackQuery,
    ReplyKeyboardMarkup, KeyboardButton,
)
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext

# --- OpenAI (async client) ---
from openai import AsyncOpenAI
# --- DOCX / PDF parsers ---
from docx import Document as DocxDocument
import fitz  # PyMuPDF

# --- .env (optional) ---
from dotenv import load_dotenv
from io import BytesIO
from datetime import datetime, timezone
from aiogram.types import BufferedInputFile

# ==================== CONFIG & LOGGING ====================
load_dotenv()

LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
)
logger = logging.getLogger("bot")

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL")  # укажите точную модель
MAX_CHARS_TO_SEND = 30_000  # мягкий лимит на длину текста
SIZE_LIMIT_MB = 20          # лимит на размер файла
UPLOADS_DIR = Path(os.getenv("UPLOADS_DIR", "uploads")).resolve()
REPORTS_DIR = Path(os.getenv("REPORTS_DIR", "reports")).resolve()

UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

if not TELEGRAM_BOT_TOKEN:
    raise RuntimeError("TELEGRAM_BOT_TOKEN не задан")
if not OPENAI_API_KEY:
    raise RuntimeError("OPENAI_API_KEY не задан")
if not OPENAI_MODEL:
    raise RuntimeError("OPENAI_MODEL не задан (например, 'gpt-5' / 'gpt-5.x')")

SYSTEM_PROMPT = '''
Ты - профессиональный судебный специалист с опытом более 10 лет.
Твоя задача — сформировать отчёт по качеству и соответствию загруженного документа на основе следующих требований.
1.	Определить тип загруженного документа (один из четырех):
Судебная экспертиза
Внесудебная экспертиза
Рецензия
Отчет об оценке
2.	Проверить наличие орфографических ошибок, вывести все.
	Пунктуацию не проверять.
3.	Проверить все нормативные документы, упомянутые в тексте:
	Определи, актуальны ли они (ФЗ, ГОСТ, СНиП, СанПиН и др.)
	Укажи, какие нормативы утратили силу.
4.	Если документ — судебная экспертиза, проверь его на соответствие 73-ФЗ “О государственной судебно-экспертной деятельности в РФ”, а также процессуальным нормам, действующим в рамках:
	ГПК РФ
	УПК РФ
	АПК РФ
	КОАП РФ
Учитывай, какой кодекс применяется в зависимости от назначения экспертизы (общая юрисдикция, арбитраж, мировые судьи и т.д.).
5.	Формат вывода:
	Установленный тобой один тип документа из четырех возможных
	Несоответствия в структуре (главы, названия)
	Найденные орфографические ошибки
	Устаревшие нормативные документы
	Если это судебная экспертиза — перечень нарушений процессуального характера
	

Дополнительные инструкции:
Не давай никаких рекомендаций, твоя задача - составить четкий детальный отчет по ТОЛЬКО имеющемуся документу.
Выводи полную информацию по каждому пунтку, не сокращай ее.
Не нужно подводить итогов и саммари в конце отчета, выполни лишь отчет по каждому пунтку.
Выведи сразу несколько переходов на новую строку, когда переходишь к следующему пункту отчета, чтобы лучше их отделить. А также разделы обозначай римскими цифрами.
Не выводи никакой информаци, кроме той, что непосредсственно относится к отчету
Если документ не поддается описанному анализу, то выведи - "Данный документ не может быть проанализирован"
'''

# ==================== INIT ====================
bot = Bot(token=TELEGRAM_BOT_TOKEN)
dp = Dispatcher()
router = Router()
oai = AsyncOpenAI(api_key=OPENAI_API_KEY)


# ==================== FSM ====================
class UploadStates(StatesGroup):
    waiting_for_file = State()


# ==================== UI HELPERS ====================
def main_kb() -> ReplyKeyboardMarkup:
    # Постоянная клавиатура, всегда видна
    return ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="Проверить документ")]],
        resize_keyboard=True, one_time_keyboard=False, selective=False
    )


# ==================== FILE/TEXT UTILS ====================
def _extract_docx_text(path: Path) -> str:
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs)

def _extract_pdf_text(path: Path) -> str:
    text_parts = []
    with fitz.open(path) as pdf:
        for page in pdf:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)

async def extract_text_from_file(path: Path, mime: Optional[str]) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx" or (mime and "officedocument.wordprocessingml.document" in mime):
        return await asyncio.to_thread(_extract_docx_text, path)
    elif suffix == ".pdf" or (mime and mime == "application/pdf"):
        return await asyncio.to_thread(_extract_pdf_text, path)
    else:
        raise ValueError("Поддерживаются только .pdf и .docx")

def looks_like_scanned(text: str) -> bool:
    letters = sum(ch.isalnum() for ch in text)
    return letters < 50

def safe_name(name: str) -> str:
    return "".join(ch for ch in name if ch.isalnum() or ch in (" ", ".", "_", "-", "(", ")")).strip() or "file"

def timestamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")


# ==================== OPENAI ====================
async def call_openai(user_text: str) -> str:
    tools = []
    tools.append({"type": "web_search_preview"})  # Встроенный веб-поиск

    t0 = time.monotonic()
    logger.info("OpenAI call (Responses): model=%s, chars=%d", OPENAI_MODEL, len(user_text))

    try:
        # В Responses API system-промпт передаётся как "instructions"
        resp = await oai.responses.create(
            model=OPENAI_MODEL,               # пример: "gpt-5"
            instructions=SYSTEM_PROMPT,       # твой большой системный промпт
            input=user_text,                  # текст документа
            tools=tools,
            max_output_tokens=20000
            # web_search + file_search
            # tool_choice="auto",             # можно явно: "auto" | {"type":"web_search_preview"} и т.п. :contentReference[oaicite:2]{index=2}
        )

        # В SDK есть удобное поле с финальным текстом
        answer = (getattr(resp, "output_text", None) or "").strip()

        # На всякий случай fallback, если модель вернула сегменты без output_text
        if not answer and getattr(resp, "output", None):
            parts = []
            for item in resp.output:
                if getattr(item, "type", "") == "message":
                    for c in getattr(item, "content", []):
                        if c.get("type") == "output_text":
                            parts.append(c.get("text", ""))
            answer = "\n".join(parts).strip()

        answer = answer.replace('/?utm_source=chatgpt.com', "")

        if not answer:
            logger.warning("OpenAI empty output_text")
            return "Пустой ответ от модели."

        dt = time.monotonic() - t0
        logger.info("OpenAI done in %.2fs", dt)
        return answer

    except Exception as e:
        logger.exception("OpenAI error: %s", e)
        raise


# ==================== REPORT BUILDER ====================
def build_report_docx(answer: str, src_filename: str | None = None, trimmed_input: bool = False) -> bytes:
    doc = DocxDocument()
    doc.core_properties.title = "Отчёт по документу"
    doc.core_properties.created = datetime.utcnow()

    doc.add_heading("Отчёт по документу", level=1)

    meta_lines = []
    if src_filename:
        meta_lines.append(f"Исходный файл: {src_filename}")
    meta_lines.append(f"Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    if trimmed_input:
        meta_lines.append("Внимание: исходный документ был обрезан по длине.")
    doc.add_paragraph(" · ".join(meta_lines))

    for block in answer.split("\n\n"):
        doc.add_paragraph(block)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ==================== HANDLERS ====================
@router.message(CommandStart())
async def start(message: Message):
    logger.info("Start command from user_id=%s username=%s",
                message.from_user.id if message.from_user else None,
                message.from_user.username if message.from_user else None)
    await message.answer(
        "Привет! Я проверю документ. Нажмите «Проверить документ» и пришлите PDF/DOCX с печатным текстом.",
        reply_markup=main_kb()
    )

# Кнопка «Проверить документ» доступна ВСЕГДА (reply keyboard)
@router.message(F.text.casefold() == "проверить документ")
async def ask_document(message: Message, state: FSMContext):
    await state.set_state(UploadStates.waiting_for_file)
    logger.info("Ask document: user_id=%s set state waiting_for_file", message.from_user.id if message.from_user else None)
    await message.answer("Пришлите файл PDF или DOCX с печатным текстом (не скан).", reply_markup=main_kb())

@router.message(UploadStates.waiting_for_file, F.document)
async def handle_document(message: Message, state: FSMContext):
    user_id = message.from_user.id if message.from_user else 0
    doc = message.document

    mime = doc.mime_type or ""
    name = (doc.file_name or "document").lower()
    logger.info("Incoming file: user_id=%s name=%s mime=%s size=%s",
                user_id, name, mime, doc.file_size)

    if not (name.endswith(".pdf") or name.endswith(".docx") or mime in (
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )):
        await message.answer("Пожалуйста, пришлите **PDF** или **DOCX** файл.", reply_markup=main_kb())
        return

    if doc.file_size and doc.file_size > SIZE_LIMIT_MB * 1024 * 1024:
        await message.answer(f"Файл слишком большой (> {SIZE_LIMIT_MB} МБ). Пришлите поменьше.", reply_markup=main_kb())
        logger.warning("File too large: user_id=%s size=%s", user_id, doc.file_size)
        return

    await message.answer("Файл получен. Читаю и отправляю в модель. Это может занять несколько минут... ⏳", reply_markup=main_kb())

    # === Сохраняем файл НА ДИСК ===
    user_dir = UPLOADS_DIR / str(user_id)
    user_dir.mkdir(parents=True, exist_ok=True)
    original_name = safe_name(doc.file_name or "upload.bin")
    saved_name = f"{timestamp()}__{original_name}"
    dest = user_dir / saved_name

    try:
        await bot.download(doc, destination=dest)
        logger.info("Saved upload: %s", dest)
    except Exception as e:
        logger.exception("Download error: %s", e)
        await message.answer("Не удалось скачать файл. Попробуйте еще раз.", reply_markup=main_kb())
        return

    # === Извлекаем текст ===
    t0 = time.monotonic()
    try:
        text = await extract_text_from_file(dest, mime)
        logger.info("Extracted text: chars=%d in %.2fs", len(text), time.monotonic() - t0)
        logger.debug("First 500 chars: %s", text[:500])
    except Exception as e:
        logger.exception("Extract error: %s", e)
        await message.answer("Не получилось прочитать файл. Убедитесь, что это PDF/DOCX с печатным текстом (не скан).", reply_markup=main_kb())
        return

    if not text.strip() or looks_like_scanned(text):
        logger.warning("Looks like scanned or empty text: user_id=%s path=%s", user_id, dest)
        await message.answer("Похоже, в файле нет извлекаемого текста (возможно, это скан). Пришлите PDF/DOCX с печатным текстом.", reply_markup=main_kb())
        return

    trimmed = False
    if len(text) > MAX_CHARS_TO_SEND:
        text = text[:MAX_CHARS_TO_SEND]
        trimmed = True
        logger.info("Trimmed input to %d chars", len(text))

    # === Вызов модели ===
    try:
        answer = await call_openai(text)
    except Exception as e:
        await message.answer("Не удалось получить ответ от модели. Попробуйте позже.", reply_markup=main_kb())
        return

    # === Сборка и отправка отчёта ===
    try:
        report_bytes = await asyncio.to_thread(
            build_report_docx,
            answer,
            src_filename=doc.file_name,
            trimmed_input=trimmed
        )
        # Сохраняем отчёт на диск (опционально, но полезно для аудита)
        reports_user_dir = REPORTS_DIR / str(user_id)
        reports_user_dir.mkdir(parents=True, exist_ok=True)
        report_path = reports_user_dir / f"{timestamp()}__report.docx"
        with open(report_path, "wb") as f:
            f.write(report_bytes)
        logger.info("Saved report: %s", report_path)

        report_file = BufferedInputFile(report_bytes, filename="report.docx")
        caption = "Готово! Отчёт во вложении."
        if trimmed:
            caption = "⚠️ Документ был очень большим — обработана только первая часть.\n" + caption

        await message.answer_document(document=report_file, caption=caption, reply_markup=main_kb())
        await state.clear()
        logger.info("Workflow finished: user_id=%s", user_id)
    except Exception as e:
        logger.exception("Report build/send error: %s", e)
        await message.answer("Не удалось сформировать Word-отчёт, отправляю текстом:\n\n" + answer, reply_markup=main_kb())

# Если ждём файл, а пришёл не файл — мягко напоминаем
@router.message(UploadStates.waiting_for_file)
async def reject_non_document(message: Message):
    logger.info("Expecting document, got message type=%s", message.content_type)
    await message.answer("Жду файл **PDF** или **DOCX**. Отправьте документ как вложение или нажмите «Проверить документ».", reply_markup=main_kb())

# Любые другие тексты вне состояния — подскажем, что делать
@router.message()
async def fallback(message: Message):
    logger.info("Fallback message from user_id=%s text=%r", message.from_user.id if message.from_user else None, message.text)
    await message.answer(
        "Нажмите «Проверить документ», затем пришлите PDF/DOCX с печатным текстом.",
        reply_markup=main_kb()
    )

# ==================== ENTRY ====================
async def main():
    dp.include_router(router)
    logger.info("Bot polling started")
    await dp.start_polling(bot)

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except (KeyboardInterrupt, SystemExit):
        logger.info("Bot stopped.")
